# -*- coding: utf-8 -*-
from __future__ import unicode_literals

from django.shortcuts import render
from django.template.response import TemplateResponse
from django.conf import settings
from django.contrib.auth.decorators import login_required
from vista.functions import *
from django.shortcuts import get_object_or_404
from vista.models import ParqueSolar, Aerogenerador
from usuarios.models import Log
from django.http import HttpResponseRedirect
from django.core.urlresolvers import reverse
from dr.forms import DRForm
from dr.models import DR, ActividadDR, ComposicionDR, FotosDR
from django.core.exceptions import PermissionDenied
import os
import StringIO
from django.http import HttpResponse
from docx import Document
from docx.shared import Mm, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from django.views.decorators.csrf import csrf_exempt
import json
from django.contrib import messages

import logging
logger = logging.getLogger('oritec')

@login_required(login_url='ingresar')
def listado(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido = ContenidoContainer()
    contenido.user = request.user
    contenido.titulo = u'Informes diarios'
    contenido.subtitulo = parque.nombre
    contenido.menu = ['menu-dr', 'menu2-listado' ]

    drs = DR.objects.filter(parque=parque)

    return TemplateResponse(request, 'dr/listado.html',
                            {'cont': contenido,
                             'parque': parque,
                             'aerogeneradores': aerogeneradores,
                             'filas': drs
                             })

@login_required(login_url='ingresar')
def agregar(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido = ContenidoContainer()
    contenido.user = request.user
    contenido.titulo = u'Informes diarios'
    contenido.subtitulo = parque.nombre
    contenido.menu = ['menu-dr', 'menu2-listado']

    form = None

    if request.method == 'POST':
        form = DRForm(request.POST)
        if form.is_valid():
            logger.debug('Formulario válido.')
            dr = form.save(commit=False)
            dr.created_by = request.user
            dr.save()
            log_msg = "Se agrega reporte diario para parque " + parque.nombre
            log = Log(texto=log_msg, tipo=1, user=request.user)
            log.save()
            #dr.save() # Necesario para actualizar campos
            messages.add_message(request, messages.SUCCESS, 'Reporte agregado')
            return HttpResponseRedirect(reverse('dr:editar', args=[parque.slug,dr.id]))
        else:
            logger.debug('Formulario no es válido...')
            messages.add_message(request, messages.ERROR, 'Error al agregar reporte')

    if form is None:
        aux = DR.objects.filter(parque=parque).order_by('-numero')
        if aux.exists():
            next_num = aux[0].numero + 1
        else:
            next_num = 1
        form = DRForm(initial={'parque':parque.id, 'numero': next_num})

    return TemplateResponse(request, 'dr/agregar.html',
                            {'cont': contenido,
                             'parque': parque,
                             'aerogeneradores': aerogeneradores,
                             'form': form,
                             'back_url': reverse('dr:listado', args=[parque.slug])
                             })

@login_required(login_url='ingresar')
def editar(request, slug, dr_id):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    dr = get_object_or_404(DR, id=dr_id)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido = ContenidoContainer()
    contenido.user = request.user
    contenido.titulo = u'Informes diarios'
    contenido.subtitulo = parque.nombre
    contenido.menu = ['menu-dr', 'menu2-listado']

    form = None

    if request.method == 'POST':
        form = DRForm(request.POST, instance=dr)
        if form.is_valid():
            logger.debug('Formulario válido.')
            dr = form.save()
            log_msg = "Se agrega reporte diario para parque " + parque.nombre
            log = Log(texto=log_msg, tipo=1, user=request.user)
            log.save()
            messages.add_message(request, messages.SUCCESS, 'Reporte Actualizado')
            #return HttpResponseRedirect(reverse('dr:observaciones-show', args=[parque.slug,observacion.id]))
        else:
            logger.debug('Formulario no es válido...')
            messages.add_message(request, messages.ERROR, 'Error al actualizar reporte')

    if form is None:
        form = DRForm(instance=dr)

    actividades = []
    nestable_id = 0
    nestable_group = 1
    for actividad in dr.actividaddr_set.all().order_by('orden'):
        nestable_id += 1
        item = { }
        if len(actividad.descripcion) > 100:
            item['descripcion']=actividad.descripcion[:100] + ' ...'
        else:
            item['descripcion'] = actividad.descripcion
        item['id'] = actividad.id
        item['nestable_id'] = nestable_id
        item['nestable_group'] = 1
        item['composiciones'] = []
        nestable_group += 1
        for composicion in actividad.composiciondr_set.all().order_by('orden'):
            comp = {}
            nestable_id += 1
            if len(composicion.pie) > 100:
                comp['pie'] = composicion.pie[:100] + ' ...'
            else:
                comp['pie'] = composicion.pie
            #comp['pie'] = composicion.pie
            comp['id'] = composicion.id
            comp['nestable_id'] = nestable_id
            comp['nestable_group'] = nestable_group

            item['composiciones'].append(comp)
        actividades.append(item)


    return TemplateResponse(request, 'dr/editar.html',
                            {'cont': contenido,
                             'parque': parque,
                             'aerogeneradores': aerogeneradores,
                             'form': form,
                             'editar': True,
                             'dr': dr,
                             'back_url': reverse('dr:listado', args=[parque.slug],),
                             'actividades':actividades
                             })

@login_required(login_url='ingresar')
def borrar(request, slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)

    if request.method == 'POST':
        logger.debug('POST del_dr')
        dr=DR.objects.get(id=int(request.POST['del_id']))
        if not (request.user.has_perm('ncr.delete_revision') or request.user == dr.created_by) :
            raise PermissionDenied

        log_msg = "Se elimina DR para parque " + parque.nombre + \
                  " - ID - " + str(dr.id)

        dr.delete()
        log = Log(texto=log_msg, tipo=3, user=request.user)
        log.save()
        messages.add_message(request, messages.SUCCESS, 'Reporte Eliminado')
        return HttpResponseRedirect(request.POST['back_url'])

@login_required(login_url='ingresar')
def actividad_agregar(request, slug, dr_id):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    dr = get_object_or_404(DR, id=dr_id)
    if request.method == 'POST':
        if 'actividad_descripcion' in request.POST:
            if request.POST['edit_actividad'] == '0':
                log_msg = "Se agrega Actividad para parque " + parque.nombre + \
                          " - DR - " + str(dr.id)
                actividad = ActividadDR(descripcion=request.POST['actividad_descripcion'],dr=dr)
                actividades = ActividadDR.objects.filter(dr=dr).order_by('-orden')
                if actividades.count() > 0:
                    actividad.orden = actividades[0].orden + 1
                actividad.save()
                log = Log(texto=log_msg, tipo=3, user=request.user)
                log.save()
                messages.add_message(request, messages.SUCCESS, 'Actividad agregada')
            else:
                actividad = ActividadDR.objects.get(id=int(request.POST['edit_actividad']))

                log_msg = "Se edita Actividad id = " + str(actividad.id) + ",para parque " + parque.nombre + \
                          " - DR - " + str(dr.id)
                actividad.descripcion = request.POST['actividad_descripcion']
                actividad.save()
                log = Log(texto=log_msg, tipo=3, user=request.user)
                log.save()
                messages.add_message(request, messages.SUCCESS, 'Actividad editada correctamente')
        return HttpResponseRedirect(reverse('dr:editar', args=[parque.slug,dr.id]))

@login_required(login_url='ingresar')
def actividad_eliminar(request, slug, dr_id):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    #dr = get_object_or_404(DR, id=dr_id)
    if request.method == 'POST':
        if 'del_id_actividad' in request.POST:
            id = int(request.POST['del_id_actividad'])
            actividad = ActividadDR.objects.get(id=id)
            actividad.delete()
            messages.add_message(request, messages.SUCCESS, 'Actividad eliminada')
        return HttpResponseRedirect(reverse('dr:editar', args=[parque.slug, dr_id]))

@login_required(login_url='ingresar')
def composicion_eliminar(request, slug, dr_id):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    #dr = get_object_or_404(DR, id=dr_id)
    if request.method == 'POST':
        if 'del_id_composicion' in request.POST:
            id = int(request.POST['del_id_composicion'])
            composicion = ComposicionDR.objects.get(id=id)
            composicion.delete()
            messages.add_message(request, messages.SUCCESS, 'Composición eliminada')
    return HttpResponseRedirect(reverse('dr:editar', args=[parque.slug, dr_id]))

@csrf_exempt
def actividad_ajax(request,slug):
    if request.is_ajax() and request.POST:
        pk = request.POST.get('id')
        info = {}
        actividad = ActividadDR.objects.get(id=pk)
        info['descripcion'] = actividad.descripcion
        info['id'] = pk
        return HttpResponse(json.dumps(info), content_type='application/json')

@csrf_exempt
def composicion_ajax(request,slug):
    if request.is_ajax() and request.POST:
        composicion_id = request.POST.get('composicion_id')
        actividad_id = request.POST.get('actividad_id')
        info = {}
        composicion = ComposicionDR.objects.get(id=composicion_id)
        info['actividad_id'] = composicion.actividad.id
        info['id'] = composicion_id
        info['pie'] = composicion.pie
        info['tipo'] = composicion.tipo
        info['fotos'] = []

        if composicion.tipo == '1V':
            info['no_fotos'] = str(1)
            info['pattern_img_id'] = str(1)
        elif composicion.tipo == '1H':
            info['no_fotos'] = str(1)
            info['pattern_img_id'] = str(2)
        elif composicion.tipo == '2V':
            info['no_fotos'] = str(2)
            info['pattern_img_id'] = str(1)
        elif composicion.tipo == '2H':
            info['no_fotos'] = str(2)
            info['pattern_img_id'] = str(2)
        elif composicion.tipo == '1V2H':
            info['no_fotos'] = str(3)
            info['pattern_img_id'] = str(1)
        elif composicion.tipo == '2H1V':
            info['no_fotos'] = str(3)
            info['pattern_img_id'] = str(2)
        elif composicion.tipo == '3V':
            info['no_fotos'] = str(3)
            info['pattern_img_id'] = str(3)
        elif composicion.tipo == '4H':
            info['no_fotos'] = str(4)
            info['pattern_img_id'] = str(1)
        elif composicion.tipo == '2V2H':
            info['no_fotos'] = str(4)
            info['pattern_img_id'] = str(2)
        elif composicion.tipo == '1V3H':
            info['no_fotos'] = str(4)
            info['pattern_img_id'] = str(3)
        elif composicion.tipo == '4V':
            info['no_fotos'] = str(4)
            info['pattern_img_id'] = str(4)
        elif composicion.tipo == '6H':
            info['no_fotos'] = str(6)
            info['pattern_img_id'] = str(1)


        for f in composicion.fotosdr_set.all():
            foto = {}
            if f.thumbnail:
                foto['url'] = f.thumbnail.url
            else:
                foto['url'] = f.imagen.url
            foto['orden'] = f.orden
            info['fotos'].append(foto)

        return HttpResponse(json.dumps(info), content_type='application/json')

@login_required(login_url='ingresar')
def composicion_agregar(request, slug, dr_id):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    dr = get_object_or_404(DR, id=dr_id)
    if request.method == 'POST':
        if 'edit_composicion' in request.POST:
            if request.POST['edit_composicion'] == '0':
                actividad = get_object_or_404(ActividadDR,id=request.POST['act_id'])
                composicion = ComposicionDR(actividad=actividad,
                                            pie=request.POST['composicion_pie'],
                                            tipo=request.POST['pattern'])
                composiciones = ComposicionDR.objects.filter(actividad=actividad).order_by('-orden')
                if composiciones.count() > 0:
                    composicion.orden = composiciones[0].orden + 1
                composicion.save()
                nofotos = int(request.POST['nofotos'])
                for i in range(nofotos):
                    file_field = 'file_input_' + request.POST['pattern'] + '_' + str(i+1)
                    if file_field in request.FILES:
                        f = FotosDR(composicion=composicion,
                                    imagen=request.FILES[file_field],
                                    orden=i)
                        f.save()
                log_msg = "Se agrega composición para actividad de parque " + parque.nombre + \
                          " - DR - " + str(dr.id)
                log = Log(texto=log_msg, tipo=3, user=request.user)
                log.save()
                messages.add_message(request, messages.SUCCESS, 'Composición agregada')
            else:
                composicion = ComposicionDR.objects.get(id=int(request.POST['edit_composicion']))
                composicion.pie=request.POST['composicion_pie']
                composicion.tipo=request.POST['pattern']
                composicion.save()
                nofotos = int(request.POST['nofotos'])
                for i in range(nofotos):
                    file_field = 'file_input_' + request.POST['pattern'] + '_' + str(i + 1)
                    if file_field in request.FILES:
                        if FotosDR.objects.filter(composicion=composicion, orden=i).exists():
                            anterior = FotosDR.objects.get(composicion=composicion, orden=i)
                            anterior.delete()
                        f = FotosDR(composicion=composicion,
                                    imagen=request.FILES[file_field],
                                    orden=i)
                        f.save()
                messages.add_message(request, messages.SUCCESS, 'Composición actualizada correctamente')

        return HttpResponseRedirect(reverse('dr:editar', args=[parque.slug,dr.id]))

def word_insert_pagebreak(document, logo_archivo, dr,codigo_informe):
    document.add_page_break()
    table = document.add_table(rows=1, cols=3, style="TablaHeader")
    row_cells = table.rows[0].cells
    r = row_cells[0].paragraphs[0].add_run()
    r.add_picture(logo_archivo, width=Mm(39), height=Mm(10.4))
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[1].paragraphs[0].add_run(codigo_informe)
    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[2].paragraphs[0].add_run('Fecha: ' + dr.fecha.strftime("%d/%m/%Y"))
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph(style='Espacio')

@login_required(login_url='ingresar')
def create_dr_word(request, slug, dr_id):
    dr = get_object_or_404(DR, id=dr_id)
    nombre_archivo = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) + '/dr/dr.docx'
    f = open(nombre_archivo, 'rb')
    document = Document(f)
    document._body.clear_content()
    temp_doc = Document()
    #document.styles.add_style('ListBullet', temp_doc.styles['ListBullet'].type)
    #document.styles.add_style('List Paragraph', temp_doc.styles['List Paragraph'].type)

    codigo_informe = 'DR_' + dr.parque.codigo.upper() + '-' + '%03d' % dr.numero

    section = document.sections[0]
    section.page_height = Mm(279.4)
    section.page_width = Mm(215.9)
    for s in temp_doc.styles.latent_styles.element:
        logger.debug(s.name)
    document.add_paragraph('')
    document.add_paragraph('')
    table = document.add_table(rows=1, cols=3, style="Titulo")

    row_cells = table.rows[0].cells
    #table.rows[0].height_rule = WD_ROW_HEIGHT.EXACTLY
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "800")
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)
    row_cells[1].paragraphs[0].add_run('INFORME DIARIO').bold = True
    logo_archivo = settings.BASE_DIR + '/static/common/images/saroenlogo-excel.png'
    r = row_cells[0].paragraphs[0].add_run()
    r.add_picture(logo_archivo, width=Cm(4.5))
    document.add_paragraph('')
    logo_empresa = dr.parque.logo.file.name
    r = row_cells[2].paragraphs[0].add_run()
    r.add_picture(logo_empresa, width=Cm(4.5))
    document.add_paragraph('')
    # Información del reporte
    table = document.add_table(rows=1, cols=4, style="Titulo")
    row_cells = table.rows[0].cells
    row_cells[0].paragraphs[0].add_run('Proyecto:').bold = True
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[1].paragraphs[0].add_run(dr.parque.nombre).bold = True
    row_cells[2].paragraphs[0].add_run('Fecha:').bold = True
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[3].paragraphs[0].add_run(dr.fecha.strftime("%d/%m/%Y")).bold = True

    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run('Cliente:').bold = True
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[1].paragraphs[0].add_run(dr.parque.cliente).bold = True
    row_cells[2].paragraphs[0].add_run('Informe Nº:').bold = True
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[3].paragraphs[0].add_run(codigo_informe).bold = True

    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run('WTG Num:').bold = True
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[1].paragraphs[0].add_run(dr.sitio)
    row_cells[2].paragraphs[0].add_run('Climatología:').bold = True
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[3].paragraphs[0].add_run(dr.climatologia)

    row_cells = table.add_row().cells
    r = row_cells[0]
    for idx in range(1,4):
        r = r.merge(row_cells[idx])

    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run('Actividades:').bold = True
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[1].paragraphs[0].add_run(dr.actividades)
    #row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = row_cells[1]
    for idx in range(2,4):
        r = r.merge(row_cells[idx])

    row_cells = table.add_row().cells
    r = row_cells[0]

    for actividad in dr.actividaddr_set.all().order_by('orden'):
        p = r.add_paragraph(actividad.descripcion)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        #p.text = 'Inspección final'
        p.style = 'ListaParrafo'

    p = r.add_paragraph()
    for idx in range(1, 4):
        r = r.merge(row_cells[idx])

    document.add_paragraph('')
    # Inserta fotos
    foto_count = 1
    for actividad in dr.actividaddr_set.all().order_by('orden'):
        for composicion in actividad.composiciondr_set.all().order_by('orden'):
            # 1 foto
            if composicion.tipo == '1H':
                table = document.add_table(rows=1, cols=1, style="TablaFotos")
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(133), height=Mm(100))
                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            elif composicion.tipo == '1V':
                table = document.add_table(rows=1, cols=1, style="TablaFotos")
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(75.2), height=Mm(100))
                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            # 2 fotos
            elif composicion.tipo == '2V':
                table = document.add_table(rows=1, cols=2, style="TablaFotos")
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(72), height=Mm(97))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(72), height=Mm(97))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            elif composicion.tipo == '2H':
                table = document.add_table(rows=1, cols=2, style="TablaFotos")
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(85), height=Mm(64))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(85), height=Mm(64))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            # 3 fotos
            elif composicion.tipo == '1V2H':
                table = document.add_table(rows=2, cols=2, style="TablaFotos")
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(72), height=Mm(97))
                cell_toadd = row_cells[0]
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(72), height=Mm(48))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # foto 3
                row_cells = table.rows[1].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=2)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(72), height=Mm(48))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell_toadd.merge(row_cells[0])
                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            elif composicion.tipo == '2H1V':
                table = document.add_table(rows=2, cols=2, style="TablaFotos")
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(72), height=Mm(48))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(72), height=Mm(97))
                cell_toadd = row_cells[1]
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # foto 3
                row_cells = table.rows[1].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=2)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(72), height=Mm(48))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell_toadd.merge(row_cells[1])
                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            elif composicion.tipo == '3V':
                table = document.add_table(rows=1, cols=3, style="TablaFotos")
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(57), height=Mm(75))
                #row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(57), height=Mm(75))
                #row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # foto 3
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[2].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(57), height=Mm(75))
                #row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            # 4 fotos
            elif composicion.tipo == '4H':
                table = document.add_table(rows=2, cols=2, style="TablaFotos")
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(70), height=Mm(52.6))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(70), height=Mm(52.6))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # foto 3
                row_cells = table.rows[1].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=2)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(70), height=Mm(52.6))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 4
                row_cells = table.rows[1].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=3)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(70), height=Mm(52.6))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            elif composicion.tipo == '2V2H':
                table = document.add_table(rows=3, cols=2, style="TablaFotos")
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(50), height=Mm(70))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell_toadd_1 = row_cells[0]
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(50), height=Mm(35))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # foto 3
                row_cells = table.rows[1].cells
                cell_toadd_1.merge(row_cells[0])

                foto = FotosDR.objects.get(composicion=composicion, orden=2)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(50), height=Mm(70))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell_toadd_2 = row_cells[1]
                # foto 4
                row_cells = table.rows[2].cells
                cell_toadd_2.merge(row_cells[1])
                foto = FotosDR.objects.get(composicion=composicion, orden=3)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(50), height=Mm(35))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            elif composicion.tipo == '1V3H':
                table = document.add_table(rows=4, cols=2, style="TablaFotos")
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(43), height=Mm(64))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row_cells[0].width = Mm(43)
                cell_toadd_1 = row_cells[0]
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(65), height=Mm(48))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[1].width = Mm(65)
                cell_toadd_2 = row_cells[1]
                # foto 3
                row_cells = table.rows[1].cells
                cell_toadd_1.merge(row_cells[0])
                cell_toadd_2.merge(row_cells[1])

                row_cells = table.rows[2].cells
                cell_toadd_1.merge(row_cells[0])

                foto = FotosDR.objects.get(composicion=composicion, orden=2)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(65), height=Mm(48))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[1].width = Mm(65)
                cell_toadd_3 = row_cells[1]
                # foto 4
                row_cells = table.rows[3].cells
                cell_toadd_3.merge(row_cells[1])
                foto = FotosDR.objects.get(composicion=composicion, orden=3)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(43), height=Mm(32))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row_cells[0].width = Mm(43)
                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
            elif composicion.tipo == '4V':
                ancho = 82
                alto = 110
                if (foto_count % 2) == 1:
                    word_insert_pagebreak(document, logo_archivo, dr,codigo_informe)
                table = document.add_table(rows=2, cols=2, style="TablaFotos")
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[0].width = Mm(ancho)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[1].width = Mm(ancho)
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # foto 3
                row_cells = table.rows[1].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=2)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[0].width = Mm(ancho)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 4
                row_cells = table.rows[1].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=3)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[1].width = Mm(ancho)
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
                # Ocupa 2 paginas
                foto_count += 1
            elif composicion.tipo == '6H':
                ancho = 85
                alto = 64
                if (foto_count % 2) == 1:
                    word_insert_pagebreak(document, logo_archivo, dr,codigo_informe)
                table = document.add_table(rows=3, cols=2, style="TablaFotos")
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells = table.rows[0].cells
                # foto 1
                foto = FotosDR.objects.get(composicion=composicion, orden=0)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[0].width = Mm(ancho)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 2
                foto = FotosDR.objects.get(composicion=composicion, orden=1)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[1].width = Mm(ancho)
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # foto 3
                row_cells = table.rows[1].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=2)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[0].width = Mm(ancho)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 4
                row_cells = table.rows[1].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=3)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[1].width = Mm(ancho)
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # foto 5
                row_cells = table.rows[2].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=4)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[0].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[0].width = Mm(ancho)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # foto 6
                row_cells = table.rows[2].cells
                foto = FotosDR.objects.get(composicion=composicion, orden=5)
                foto_filename = settings.BASE_DIR + foto.imagen.url
                r = row_cells[1].paragraphs[0].add_run()
                r.add_picture(foto_filename, width=Mm(ancho), height=Mm(alto))
                row_cells[1].width = Mm(ancho)
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                p = document.add_paragraph('Ilustración ' + str(foto_count) + '. ' + composicion.pie)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = 'Pie'
                # Ocupa 2 paginas
                foto_count += 1

            foto_count += 1
            if (foto_count % 2) == 0:
                word_insert_pagebreak(document, logo_archivo, dr, codigo_informe)
            else:
                document.add_paragraph()

    #core_properties = document.core_properties
    #core_properties.title =
    #core_properties.comments = 'Fecha: ' + dr.fecha.strftime("%d/%m/%Y")

    target_stream = StringIO.StringIO()
    document.save(target_stream)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    fecha_str = dr.fecha.strftime("%d%m%y")
    nombre = codigo_informe + '_' + fecha_str + '.docx'
    response['Content-Disposition'] = 'attachment; filename=' + nombre
    target_stream.flush()
    ret_word = target_stream.getvalue()
    target_stream.close()
    response.write(ret_word)
    return response

@csrf_exempt
def dr_ordenar(request, slug, dr_id):
    if request.is_ajax() and request.POST:
        parque = get_object_or_404(ParqueSolar, slug=slug)
        dr = get_object_or_404(DR, id=dr_id)
        orden = json.loads(request.POST.get('orden'))
        orden_actividad = 0
        for act in orden:
            actividad = ActividadDR.objects.get(id=act['id'])
            actividad.orden = orden_actividad
            actividad.save()
            orden_composicion = 0
            if 'children' in act:
                for comp in act['children']:
                    composicion = ComposicionDR.objects.get(id=comp['id'])
                    composicion.actividad = actividad
                    composicion.orden = orden_composicion
                    composicion.save()
                    orden_composicion += 1
            orden_actividad += 0
        return HttpResponse(json.dumps('OK'), content_type='application/json')