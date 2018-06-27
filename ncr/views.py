# -*- coding: utf-8 -*-
import sys
reload(sys)
sys.setdefaultencoding("utf-8")
from django.template.response import TemplateResponse
from vista.functions import *
from django.contrib.auth.decorators import login_required, permission_required
from vista.models import ParqueSolar,Aerogenerador
from django.http import HttpResponse, HttpResponseRedirect
from django.core.urlresolvers import reverse
from django.shortcuts import get_object_or_404
from forms import ObservacionForm, RevisionForm, RevisionFormFull, NCR, Punchlist,DuplicarObservacionForm
from ncr.models import Observacion, Revision, Fotos, Observador,Componente,Subcomponente,Tipo,Severidad,EstadoRevision,Prioridad
from easy_pdf.rendering import render_to_pdf_response, render_to_pdf
from vista.models import Aerogenerador
import json
import logging
import os
import urlparse
import StringIO
import zipfile
import base64
from docx import Document
from docx.shared import Mm, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
from openpyxl import Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
from datetime import date
from datetime import datetime
from django.core.exceptions import PermissionDenied
from usuarios.models import Log
from django.db import transaction
from django.contrib import messages

from slimit import ast
from slimit.parser import Parser
from slimit.visitors import nodevisitor

import requests

from django.db import IntegrityError

logger = logging.getLogger('oritec')

@login_required(login_url='ingresar')
def home(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Parque Eólico-NCR'
    contenido.subtitulo=parque.nombre
    contenido.menu = ['menu-principal', 'menu2-resumen']

    return TemplateResponse(request, 'vista/home.html',
        {'cont': contenido,
         'parque': parque,
         'aerogeneradores':aerogeneradores,
        })

def serializeGrafico(d):
    s = '['
    if len(d) > 0:
        for v in d:
            s += '{'
            for key, value in v.iteritems():
                if isinstance(value, (int, long, float, complex)):
                    if isinstance(value, (float)):
                        s += key + ':'+ "%.1f" % value + ','
                    else:
                        s += key + ':' + str(value) + ','
                elif isinstance(value,(dict,list)):
                    s2 = serializeGrafico(value)
                    s += key + ':' + s2 + ','
                else:
                    s +=  key + ':"' + value + '",'

            s = s[:-1]
            s += '},'
        s = s[:-1]

    s += ']'

    return s

def graficoBarrasSimple(resultados,field,secciones, showall = False):
    data_graficos = []
    count = 0
    for s in secciones:
        karws = {field: s}
        aux=resultados.filter(**karws)
        if (aux.count()>0 or showall):
            data_graficos.append({"name":s.graphText(),"y":aux.count()})
            count += 1

    datos= serializeGrafico(data_graficos)
    return datos

def graficoBarrasCompleto(resultados, field, secciones, showall = False):
    data_full = []
    count = 0
    filtros = Severidad.objects.all()
    for f in filtros:
        data_graficos = []
        karws = {'severidad': f}
        resultados_filtrados = resultados.filter(**karws)
        for s in secciones:
            karws = {field: s}
            aux=resultados_filtrados.filter(**karws)
            if (aux.count()>0 or showall):
                data_graficos.append({"name":s.graphText(),"y":aux.count()})
                count += 1
        data_full.append({"name": f.graphText(), "data": data_graficos})
    datos = serializeGrafico(data_full)
    return datos

@login_required(login_url='ingresar')
def observaciones_resumen(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Resumen de Observaciones'
    contenido.subtitulo='Parque '+parque.nombre
    contenido.menu = ['menu-ncr', 'menu2-observaciones-resumen']
    observaciones = Observacion.objects.filter(parque=parque)
    url_append = ''
    table_show_ag = True

    observaciones2 = observaciones.exclude(cerrado = True)

    grafico_estado = graficoBarrasSimple(observaciones2,'estado',EstadoRevision.objects.all().order_by('-id'),showall=True)

    observaciones2 = observaciones2.exclude(estado__nombre__exact='Solucionado')

    grafico_severidad = graficoBarrasSimple(observaciones2,'severidad',Severidad.objects.all(), showall=True)
    grafico_prioridad = graficoBarrasSimple(observaciones2, 'prioridad', Prioridad.objects.all(), showall=True)
    grafico_componente = graficoBarrasSimple(observaciones2,'componente',Componente.objects.all())
    grafico_subcomponente = graficoBarrasSimple(observaciones2, 'sub_componente', Subcomponente.objects.all())
    grafico_tipo = graficoBarrasSimple(observaciones2, 'tipo', Tipo.objects.all())

    grafico_aerogenerador = graficoBarrasCompleto(observaciones2,'aerogenerador',aerogeneradores, showall=True)

    return TemplateResponse(request, 'ncr/resumen.html',
                  {'cont': contenido,
                   'parque': parque,
                   'observaciones': observaciones,
                   'url_append':url_append,
                   'table_show_ag': table_show_ag,
                   'aerogeneradores':aerogeneradores,
                   'grafico_estado':grafico_estado,
                   'grafico_severidad': grafico_severidad,
                   'grafico_prioridad': grafico_prioridad,
                   'grafico_componente': grafico_componente,
                   'grafico_subcomponente': grafico_subcomponente,
                   'grafico_tipo': grafico_tipo,
                   'grafico_aerogenerador': grafico_aerogenerador,
                   })

@login_required(login_url='ingresar')
def observaciones_duplicadas(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Observaciones Duplicadas'
    contenido.subtitulo='Parque '+parque.nombre
    contenido.menu = ['menu-ncr', 'menu2-duplicadas']
    observaciones = Observacion.objects.filter(parque=parque, copied=True)
    url_append = ''
    table_show_ag = True

    if request.method == 'POST':
        obs = Observacion.objects.get(id=int(request.POST['remover_id']))
        obs.copied = False
        obs.save()
        messages.add_message(request, messages.SUCCESS, 'Observación removida del listado')

    return TemplateResponse(request, 'ncr/duplicadas.html',
                  {'cont': contenido,
                   'parque': parque,
                   'observaciones': observaciones,
                   'url_append':url_append,
                   'table_show_ag': table_show_ag,
                   'aerogeneradores':aerogeneradores,
                   })

@login_required(login_url='ingresar')
def observaciones(request,slug,slug_ag):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    aerogenerador = get_object_or_404(Aerogenerador,parque=parque,slug=slug_ag)
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'NCR Aerogenerador '+ aerogenerador.nombre
    contenido.subtitulo='Parque '+ parque.nombre
    contenido.menu = ['menu-ncr', 'menu2-observaciones-'+str(aerogenerador.idx)]
    url_append='?aerogenerador='+str(aerogenerador.idx)

    observaciones = Observacion.objects.filter(aerogenerador__idx__exact=aerogenerador.idx, parque= parque)

    observaciones2 = observaciones.exclude(cerrado=True)
    grafico_estado = graficoBarrasSimple(observaciones2, 'estado', EstadoRevision.objects.all().order_by('-id'),
                                         showall=True)
    observaciones2 = observaciones.exclude(estado__nombre__exact='Solucionado')
    grafico_severidad = graficoBarrasSimple(observaciones2, 'severidad', Severidad.objects.all(), showall=True)
    grafico_prioridad = graficoBarrasSimple(observaciones2, 'prioridad', Prioridad.objects.all(), showall=True)
    grafico_componente = graficoBarrasSimple(observaciones2, 'componente', Componente.objects.all())
    grafico_subcomponente = graficoBarrasSimple(observaciones2, 'sub_componente', Subcomponente.objects.all())
    grafico_tipo = graficoBarrasSimple(observaciones2, 'tipo', Tipo.objects.all())

    return TemplateResponse(request, 'ncr/resumen.html',
        {'cont': contenido,
            'parque': parque,
            'observaciones': observaciones,
            'url_append': url_append,
            'aerogeneradores':aerogeneradores,
            'grafico_estado':grafico_estado,
            'grafico_severidad': grafico_severidad,
            'grafico_prioridad': grafico_prioridad,
            'grafico_componente': grafico_componente,
            'grafico_subcomponente': grafico_subcomponente,
            'grafico_tipo': grafico_tipo,
            'wtg': aerogenerador
        })

@login_required(login_url='ingresar')
@permission_required('ncr.add_observacion', raise_exception=True)
def add_observacion(request,slug,observacion_id=0):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    edit_observacion = None
    if observacion_id != 0:
        edit_observacion = get_object_or_404(Observacion,id=observacion_id)
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.subtitulo=parque.nombre
    if 'aerogenerador' in request.GET:
        contenido.titulo = u'Agregar Observación'
        contenido.menu = ['menu-ncr', 'menu2-observaciones-'+str(request.GET['aerogenerador'])]
        ag_readonly = True
        back_url = reverse('ncr:observaciones', args=[parque.slug,request.GET['aerogenerador']])
    elif edit_observacion is not None:
        if not (request.user.has_perm('ncr.change_observacion') or request.user == edit_observacion.created_by) :
            raise PermissionDenied
        contenido.titulo = u'Editar Observación'
        contenido.menu = ['menu-ncr', 'menu2-observaciones-' + str(edit_observacion.aerogenerador.idx)]
        ag_readonly = False
        back_url = reverse('ncr:observaciones-show', args=[parque.slug, str(edit_observacion.id)])
    else:
        contenido.titulo = u'Agregar Observación'
        contenido.menu = ['menu-ncr', 'menu2-observaciones-resumen']
        ag_readonly = False
        back_url = reverse('ncr:observaciones-resumen', args=[parque.slug])

    observacionForm = None

    #logger.debug(request.GET['aerogenerador'])

    if request.method == 'POST':
        if edit_observacion is not None:
            if not request.user.has_perm('ncr.change_observacion'):
                raise PermissionDenied
            observacionForm = ObservacionForm(request.POST,instance=edit_observacion)
            rev = edit_observacion.revision_set.order_by('id')[0]
            revisionForm = RevisionForm(request.POST,instance=rev)
        else:
            observacionForm = ObservacionForm(request.POST,initial={"parque":parque})
            revisionForm = RevisionForm(request.POST)
        if observacionForm.is_valid() and revisionForm.is_valid():
            if edit_observacion is not None:
                observacion = observacionForm.save()
                revision = revisionForm.save()
                # es necesario para actualizar el estado de la observacion
                observacion.save()
                log_msg = 'Se edita observación para parque ' + parque.nombre + \
                          " - Aerogenerador - " + observacion.aerogenerador.nombre + \
                          " - " + observacion.nombre
                log = Log(texto=log_msg, tipo=2, user=request.user)
                log.save()
            else:
                observacion = observacionForm.save(commit=False)
                observacion.created_by = request.user
                observacion.save()
                revision = revisionForm.save(commit=False)
                revision.observacion = observacion
                revision.created_by = request.user
                revision.reported_by = observacion.reported_by
                revision.save()
                # es necesario para actualizar el estado de la observacion

                log_msg = 'Se crea observación para parque ' + parque.nombre + \
                          " - Aerogenerador - " + observacion.aerogenerador.nombre + \
                          " - " + observacion.nombre
                log = Log(texto=log_msg, tipo=1, user=request.user)
                log.save()
                observacion.save()
            return HttpResponseRedirect(reverse('ncr:observaciones-show', args=[parque.slug,observacion.id]))
        else:
            logger.debug('Formulario no es válido...')

    if observacionForm is None:
        if edit_observacion is not None:
            observacionForm = ObservacionForm(instance=edit_observacion)
            rev = edit_observacion.revision_set.order_by('id')[0]
            revisionForm = RevisionForm(instance=rev)
        else:
            if 'aerogenerador' in request.GET:
                ag = get_object_or_404(Aerogenerador, idx=int(request.GET['aerogenerador']), parque=parque)
                observacionForm = ObservacionForm(initial={"parque":parque,"aerogenerador":ag.id})
            else:
                observacionForm = ObservacionForm(initial={"parque": parque})
            revisionForm = RevisionForm()

    return TemplateResponse(request, 'ncr/agregarObservacion.html',
        {'cont': contenido,
         'parque': parque,
         'observacionForm':observacionForm,
         'revisionForm': revisionForm,
         'ag_readonly': ag_readonly,
         'back_url': back_url,
         'aerogeneradores':aerogeneradores,
         'edit_observacion': edit_observacion,
        })

@login_required(login_url='ingresar')
def show_observacion(request,slug,observacion_id):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    observacion=get_object_or_404(Observacion, id=observacion_id)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Observación OBS_' + parque.codigo +\
                     '-' + observacion.aerogenerador.nombre + '-' + str(observacion.observacion_id)
    contenido.subtitulo=parque.nombre
    contenido.menu = ['menu-ncr', 'menu2-observaciones-'+str(observacion.aerogenerador.idx)]
    main_fotos = {}
    fotos = {}
    for r in observacion.revision_set.all().order_by('id'):
        #logger.debug(r)
        fotos[r.id]=[]
        results = Fotos.objects.filter(revision=r,principal=True).order_by('orden')
        if results.count() > 0:
            main_fotos[r.id]=results[0].thumbnail.url
        results = Fotos.objects.filter(revision=r)
        for foto in results:
            fotos[r.id].append(foto.imagen.url)

    template_name = 'ncr/showObservacion.html'
    duplicarForm = None
    if request.method == 'POST':
        if 'duplicar' in request.POST:
            duplicarForm = DuplicarObservacionForm(request.POST, parque=parque)
            if duplicarForm.is_valid():
                ags = duplicarForm.cleaned_data['aerogenerador']
                for ag in ags:
                    obs_pk = observacion.pk
                    new_obs = observacion
                    new_obs.pk = None
                    new_obs.copied = True
                    new_obs.aerogenerador = ag
                    new_obs.save()
                    observacion = Observacion.objects.get(pk=obs_pk)
                    for rev in observacion.revision_set.all():
                        rev_pk = rev.pk
                        new_rev = rev
                        new_rev.pk = None
                        new_rev.observacion = new_obs
                        new_rev.save()
                        current_rev = Revision.objects.get(pk=rev_pk)
                        for foto in current_rev.fotos_set.all():
                            new_foto = foto
                            new_foto.pk = None
                            new_foto.revision = new_rev
                            new_foto.save()
                messages.add_message(request, messages.SUCCESS, 'Observación duplicada')
            else:
                messages.add_message(request, messages.ERROR, 'Problema al duplicar observación')

    if duplicarForm is None:
        duplicarForm = DuplicarObservacionForm(parque=parque)

    if 'printable' in request.GET:
        template_name = 'ncr/showObservacion-printable.html'
        breadcrumbs = False
    else:
        breadcrumbs = True

    return TemplateResponse(request, template_name,
        {'cont': contenido,
         'parque': parque,
         'observacion': observacion,
         'main_fotos': main_fotos,
         'fotos': fotos,
         'aerogeneradores':aerogeneradores,
         'breadcrumbs': breadcrumbs,
         'duplicarForm': duplicarForm
        })

@login_required(login_url='ingresar')
@permission_required('ncr.delete_observacion', raise_exception=True)
def del_observacion(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Listado de observadores'
    contenido.subtitulo='Parque '+ parque.nombre
    contenido.menu = ['menu-principal', 'menu2-observadores']
    observadores = Observador.objects.all().order_by('id')
    response_data = {}

    if request.method == 'POST':
        logger.debug('POST del_observacion')
        obs = Observacion.objects.get(id=int(request.POST['del_id']))
        log_msg = "Se elimina observación para parque " + parque.nombre + \
                  " - Aerogenerador - " + obs.aerogenerador.nombre + \
                  " - Nombre" + obs.nombre
        log = Log(texto=log_msg, tipo=3, user=request.user)
        log.save()
        if not (request.user.has_perm('ncr.delete_observacion') or request.user == obs.created_by) :
            raise PermissionDenied
        obs.delete()
        return HttpResponseRedirect(request.POST['back_url'])

@login_required(login_url='ingresar')
def add_images(request,slug,revision_id):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    revision= Revision.objects.get(pk=revision_id)
    logger.debug('Enter add_images')
    response_data = {}
    response_data['files'] = []
    if not (request.user.has_perm('ncr.change_revision') or request.user == revision.created_by):
        raise PermissionDenied

    results = Fotos.objects.filter(revision=revision).order_by('orden')
    if results.count() > 0:
        last_orden = results.reverse()[0].orden
    else:
        last_orden = 0

    if request.method == 'POST':
        for key, file in request.FILES.items():
            #logger.debug(file)
            primary = False
            with transaction.atomic():
                results = Fotos.objects.filter(revision=revision).order_by('orden')
                if results.count() > 0:
                    last_orden = results.reverse()[0].orden
                else:
                    last_orden = 0
                instance = Fotos(imagen=file,revision=revision,orden=last_orden+1)
                last_orden += 1
                instance.save()
            logger.debug('Imagen:'+ file.name + ', orden=' + str(instance.orden))
            log_msg = "Se agrega imagen para parque " + parque.nombre + \
                      " - Aerogenerador " + revision.observacion.aerogenerador.nombre + \
                      " - Observacion " + revision.observacion.nombre + \
                      " - Nombre " + file.name
            log = Log(texto=log_msg, tipo=1, user=request.user)
            log.save()
            principales =request.POST.getlist('radio')
            if file.name in principales:
                set_primary_photo(foto_id=instance.id)
                primary = True
                log_msg = "Se agrega imagen principal para parque " + parque.nombre + \
                          " - Aerogenerador " + revision.observacion.aerogenerador.nombre + \
                          " - Observacion " + revision.observacion.nombre + \
                          " - Nombre " + file.name
                log = Log(texto=log_msg, tipo=1, user=request.user)
                log.save()

            data={'name': file.name,
                  'size':str(file.size),
                  'url':instance.imagen.url,
                  'thumbnailUrl':instance.thumbnail.url,
                  'deleteUrl':reverse('ncr:imagenes-delete',args=[parque.slug,instance.id]),
                  "deleteType": "DELETE",
                  "mainPhoto": primary,
                  "photoId" : str(instance.id)}
            response_data['files'].append(data)
            #logger.debug(instance.imagen.url)
        response=json.dumps(response_data)

        return HttpResponse(
            response,
            content_type="application/json"
        )

@login_required(login_url='ingresar')
def list_fotos(request,slug):
    logger.debug("Enter list_fotos")
    parque = get_object_or_404(ParqueSolar, slug=slug)
    response_data = {}
    response_data['files'] = []
    if request.method == 'POST':
        revision = Revision.objects.get(pk=request.POST['revision_id'])
        results = Fotos.objects.filter(revision=revision).order_by('orden')
        for foto in results:
            data={'name': os.path.basename(foto.imagen.name),
                  'size':str(foto.imagen.size),
                  'url':foto.imagen.url,
                  'thumbnailUrl':foto.thumbnail.url,
                  'deleteUrl':reverse('ncr:imagenes-delete',args=[parque.slug,foto.id]),
                  "deleteType": "DELETE",
                  "mainPhoto": foto.principal,
                  "photoId" : str(foto.id)}
            response_data['files'].append(data)

        response=json.dumps(response_data)

        return HttpResponse(
            response,
            content_type="application/json"
        )

@login_required(login_url='ingresar')
def table_fotos(request,slug,revision_id):
    logger.debug("Enter list_fotos")
    parque = get_object_or_404(ParqueSolar, slug=slug)
    response_data = {}
    response_data['data'] = []
    revision = Revision.objects.get(pk=revision_id)
    results = Fotos.objects.filter(revision=revision).order_by('orden')
    if results.count()> 0:
        last_orden = results.reverse()[0].orden
    else:
        last_orden = 0
    fotos = []
    for foto in results:
        if foto.orden == 0:
            foto.orden = last_orden + 1
            last_orden += 1
            fotos.append(foto)
        data = [foto.orden,
                foto.thumbnail.url,
                foto.principal,
                foto.id
                ]
        response_data['data'].append(data)
    for foto in fotos:
        foto.save(update_fields=['orden'])
    response = json.dumps(response_data)
    return HttpResponse(
        response,
        content_type="application/json"
    )

@login_required(login_url='ingresar')
def del_image(request,slug,image_id):
    logger.debug("Enter del_image")
    response_data = {}
    response_data['files'] = []
    if request.method == 'DELETE':
        foto = Fotos.objects.get(pk=image_id)

        log_msg = "Se elimina imagen "+ foto.imagen.name
        log = Log(texto=log_msg, tipo=3, user=request.user)
        log.save()
        curr_orden = foto.orden
        revision = foto.revision
        foto.delete()
        siguientes = Fotos.objects.filter(revision=revision,orden__gt=curr_orden).order_by('orden')
        lista = []
        for siguiente in siguientes:
            siguiente.orden = curr_orden
            curr_orden += 1
            lista.append(siguiente)
        for l in lista:
            l.save()
        data = {os.path.basename(foto.imagen.name):True}
        response_data['files'].append(data)
        response=json.dumps(response_data)

        return HttpResponse(
            response,
            content_type="application/json"
        )

def set_primary_photo(foto_id, estado = True):
    foto = Fotos.objects.get(pk=foto_id)
    #logger.debug("Enter set_primary_photo: " + foto.imagen.name + ', estado = '+primary_image)
    foto.principal = estado
    foto.save(update_fields=["principal"])

@login_required(login_url='ingresar')
def primary_image(request,slug):
    if request.method == 'POST':
        #logger.debug("primary image")
        if request.POST['estado'] == 'true':
            estado = True
        else:
            estado = False
        set_primary_photo(foto_id=request.POST['foto_id'], estado = estado)
        foto = Fotos.objects.get(pk=request.POST['foto_id'])
        logger.debug('foto actualizada')
        if estado:
            log_msg = "Se escoge imagen hacia reporte " + foto.imagen.name
        else:
            log_msg = "Se extrae imagen de reporte " + foto.imagen.name
        log = Log(texto=log_msg, tipo=2, user=request.user)
        log.save()
        fotos = Fotos.objects.filter(principal=True,revision=foto.revision).order_by('orden')
        if fotos.count() > 0:
            foto2 = fotos[0]
        else:
            foto2 = Fotos.objects.filter(revision=foto.revision).order_by('orden')[0]
        return HttpResponse(foto2.thumbnail.url)

@login_required(login_url='ingresar')
def set_orden(request,slug):
    if request.method == 'POST':
        #logger.debug("primary image")
        revision_id = int(request.POST['revision_id'])
        old = request.POST.getlist('old[]')
        new = request.POST.getlist('new[]')
        fotos = []
        for idx,val in enumerate(old):
            foto = Fotos.objects.get(revision_id=revision_id,orden=int(val))
            foto.orden = int(new[idx])
            fotos.append(foto)
        for foto in fotos:
            foto.save(update_fields=['orden'])

        fotos = Fotos.objects.filter(principal=True,revision_id=revision_id).order_by('orden')
        if fotos.count() > 0:
            return HttpResponse(fotos[0].thumbnail.url)
        else:
            return HttpResponse(Fotos.objects.filter(revision_id=revision_id).order_by('orden')[0].thumbnail.url)
        #return HttpResponse('ok')
    else:
        return HttpResponse('fail')

@login_required(login_url='ingresar')
def add_revision(request,slug,observacion_id, revision_id = 0):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    edit_revision = None
    if revision_id != 0:
        edit_revision = get_object_or_404(Revision, id=revision_id)
    contenido=ContenidoContainer()
    contenido.user=request.user
    if edit_revision is not None:
        contenido.titulo = u'Editar Revisión'
        if not (request.user.has_perm('ncr.change_revision') or request.user == edit_revision.created_by) :
            raise PermissionDenied
    else:
        contenido.titulo=u'Agregar Revisión'
    contenido.subtitulo=parque.nombre
    observacion = get_object_or_404(Observacion, id=observacion_id)
    contenido.menu = ['menu-ncr', 'menu2-observaciones-' + str(observacion.aerogenerador.idx)]

    revisionForm = None

    if request.method == 'POST':
        if edit_revision is not None:
            revisionForm = RevisionFormFull(request.POST, instance=edit_revision)
        else:
            revisionForm = RevisionFormFull(request.POST)
        if revisionForm.is_valid():
            logger.debug('Formulario válido.')
            if edit_revision is not None:
                revision = revisionForm.save()
                log_msg = "Se edita revisión para parque " + parque.nombre + \
                          " - Aerogenerador - " + revision.observacion.aerogenerador.nombre + \
                          " - Observacion " + revision.observacion.nombre
                log = Log(texto=log_msg, tipo=2, user=request.user)
                log.save()
            else:
                revision = revisionForm.save(commit=False)
                revision.created_by = request.user
                revision.save()
                log_msg = "Se agrega revisión para parque " + parque.nombre + \
                          " - Aerogenerador - " + revision.observacion.aerogenerador.nombre + \
                          " - Observacion " + revision.observacion.nombre
                log = Log(texto=log_msg, tipo=1, user=request.user)
                log.save()
            revision.observacion.save() # Necesario para actualizar campos
            return HttpResponseRedirect(reverse('ncr:observaciones-show', args=[parque.slug,observacion.id]))
        else:
            logger.debug('Formulario no es válido...')

    if revisionForm is None:
        if edit_revision is not None:
            revisionForm = RevisionFormFull(instance=edit_revision)
        else:
            initial = {"observacion":observacion,
                       "nombre":observacion.nombre,
                       "prioridad": observacion.revision_set.all().order_by('-id')[0].prioridad}
            revisionForm = RevisionFormFull(initial=initial)

    return TemplateResponse(request, 'ncr/agregarRevision.html',
        {'cont': contenido,
         'parque': parque,
         'observacion': observacion,
         'revisionForm': revisionForm,
         'aerogeneradores':aerogeneradores,
         'edit_revision': edit_revision,
        })

def generateExcelNCR(resultados):
    wb = Workbook()
    ws = wb.active
    ws.title = "ReporteNCR"
    # Titulos
    ws['A1'] = '#'
    ws['B1'] = 'WTG'
    ws['C1'] = 'Estado'
    ws['D1'] = 'Severidad'
    ws['E1'] = 'Prioridad'
    ws['F1'] = 'Componente'
    ws['G1'] = 'Subcomponente'
    ws['H1'] = 'Tipo'
    ws['I1'] = 'Descripcion'
    row = 1
    for r in resultados:
        row += 1
        ws.cell(row=row, column=1, value=str(row-1))
        ws.cell(row=row, column=2,value=r.aerogenerador.nombre)
        ws.cell(row=row, column=3, value=r.estado.nombre)
        ws.cell(row=row, column=4, value=r.severidad.nombre)
        ws.cell(row=row, column=5, value=r.prioridad.nombre)
        ws.cell(row=row, column=6, value=r.componente.nombre)
        ws.cell(row=row, column=7, value=r.sub_componente.nombre)
        ws.cell(row=row, column=8, value=r.tipo.nombre)
        ws.cell(row=row, column=9, value=r.nombre)

    tab = Table(displayName="NCR", ref="A1:I"+str(row))
    style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False,
                           showLastColumn=False, showRowStripes=True, showColumnStripes=False)
    tab.tableStyleInfo = style
    ws.add_table(tab)
    dims = {}
    for row in ws.rows:
        for cell in row:
            if cell.value:
                dims[cell.column] = max((dims.get(cell.column, 0), len(cell.value)+4))
    for col, value in dims.items():
        ws.column_dimensions[col].width = value
    target_stream = StringIO.StringIO()
    wb.save(target_stream)
    return target_stream

@login_required(login_url='ingresar')
def informeNCR(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Informe NCRs'
    contenido.subtitulo='Parque '+ parque.nombre
    contenido.menu = ['menu-ncr', 'menu2-informeNCR']

    form = NCR(parque=parque)
    resultados = None

    grafico_estado = ''
    grafico_severidad = ''
    grafico_componente = ''
    grafico_subcomponente = ''
    grafico_tipo = ''

    graficos = []

    if request.method == 'POST':
        logger.debug('informeNCR Post')
        form = NCR(request.POST,parque=parque)
        if form.is_valid():
            logger.debug("Form Valid")
            resultados = Observacion.objects.filter(aerogenerador__in=form.cleaned_data['aerogenerador'])
            resultados = resultados.exclude(cerrado = True)
            resultados = resultados.filter(estado__in=form.cleaned_data['estado'])
            resultados = resultados.filter(severidad__in=form.cleaned_data['severidad'])
            resultados = resultados.filter(prioridad__in=form.cleaned_data['prioridad'])
            resultados = resultados.filter(componente__in=form.cleaned_data['componente'])
            resultados = resultados.filter(sub_componente__in=form.cleaned_data['subcomponente'])
            resultados = resultados.filter(tipo__in=form.cleaned_data['tipo'])
            resultados = resultados.filter(fecha_observacion__gte=form.cleaned_data['fecha_from'],
                                           fecha_observacion__lte=form.cleaned_data['fecha_to'])
            if (len(form.cleaned_data['condicion']) == 1):
                if 'reparadas' in form.cleaned_data['condicion']:
                    condicion = True
                else:
                    condicion = False
                resultados = resultados.filter(cerrado = condicion)
            elif (len(form.cleaned_data['condicion']) == 0):
                resultados = []

            if (len(form.cleaned_data['clase']) == 1):
                if '1' in form.cleaned_data['clase']:
                    clase = True
                else:
                    clase = False
                resultados = resultados.filter(clase = clase)
            elif (len(form.cleaned_data['clase']) == 0):
                resultados = []

            if (len(form.cleaned_data['punchlist']) == 1):
                if '1' in form.cleaned_data['punchlist']:
                    punchlist = True
                else:
                    punchlist = False
                resultados = resultados.filter(punchlist = punchlist)
            elif (len(form.cleaned_data['punchlist']) == 0):
                resultados = []
            #for key,value in form.cleaned_data.iteritems():
            #    logger.debug(key)
            logger.debug(resultados)

            if len(form.cleaned_data['estado']) > 1:
                grafico_estado = graficoBarrasSimple(resultados, 'estado', EstadoRevision.objects.all().order_by('-id'),
                                                 showall=True)

            if len(form.cleaned_data['severidad']) > 1:
                grafico_severidad = graficoBarrasSimple(resultados, 'severidad', Severidad.objects.all(), showall=True)
                graficos.append('severidad')
            if len(form.cleaned_data['componente']) > 1:
                grafico_componente = graficoBarrasSimple(resultados, 'componente', Componente.objects.all())
                graficos.append('componente')
            if len(form.cleaned_data['subcomponente']) > 1:
                grafico_subcomponente = graficoBarrasSimple(resultados, 'sub_componente', Subcomponente.objects.all())
                graficos.append('subcomponente')
            if len(form.cleaned_data['tipo']) > 1:
                grafico_tipo = graficoBarrasSimple(resultados, 'tipo', Tipo.objects.all())
                graficos.append('tipo')

            if 'excel' in request.POST:
                if not request.user.has_perm('usuarios.create_editables') :
                    raise PermissionDenied
                logger.debug('Excel')
                resultados = resultados.order_by('aerogenerador__idx')
                target_stream = generateExcelNCR(resultados)
                response = HttpResponse(
                    content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                response['Content-Disposition'] = 'attachment; filename="ReporteNCR.xlsx"'
                target_stream.flush()
                ret_excel = target_stream.getvalue()
                target_stream.close()
                response.write(ret_excel)
                return response
            if 'pdf' in request.POST:
                logger.debug('PDF')
                resultados = resultados.order_by('aerogenerador__idx')
                imagenes = listFotos_v2(resultados)
                if "colores" in request.POST:
                    colores = True
                else:
                    colores = False
                if "estados" in request.POST:
                    estados = True
                else:
                    estados = False
                fecha = datetime.strptime(request.POST['fecha'],'%d-%m-%Y').date()

                nombre_archivo = 'Sin'
                if request.POST['nombre'] != '':
                    nombre_archivo = request.POST['nombre']
                nombre = 'INFNCR_' + parque.codigo + '-' + nombre_archivo + '_' + fecha.strftime("%y%m%d") + '.pdf'

                respuesta = generatePdf(parque,resultados,imagenes,request.POST['titulo'],request,
                                        colores=colores,
                                        estados = estados,
                                        fecha=fecha,
                                        nombre = nombre)
                respuesta['Content-Disposition'] = 'attachment; filename="' + nombre + '"'
                return respuesta

    return TemplateResponse(request, 'ncr/informeNCR.html',
        {'cont': contenido,
         'parque': parque,
         'aerogeneradores':aerogeneradores,
         'form':form,
         'resultados':resultados,
         'grafico_estado':grafico_estado,
         'grafico_severidad': grafico_severidad,
         'grafico_componente': grafico_componente,
         'grafico_subcomponente': grafico_subcomponente,
         'grafico_tipo': grafico_tipo,
         'graficos':graficos,
        })

def generatePdf(parque,resultados,imagenes, titulo,
                request = None,
                show_fotos=True,
                colores = True,
                estados = True,
                fecha = date.today,
                nombre = ''):

    with open(os.path.join(settings.BASE_DIR, 'static/common/images/saroenlogo.png'), "rb") as image_file:
        logo_saroen = base64.b64encode(image_file.read())
    with open(os.path.join(settings.BASE_DIR, 'static/common/images/check-mark-3-64.gif'), "rb") as image_file:
        img_solucionado = base64.b64encode(image_file.read())
    with open(os.path.join(settings.BASE_DIR, 'static/common/images/x-mark-64-amarillo.gif'), "rb") as image_file:
        img_parcialsolucionado = base64.b64encode(image_file.read())
    with open(os.path.join(settings.BASE_DIR, 'static/common/images/x-mark-64.gif'), "rb") as image_file:
        img_nosolucionado = base64.b64encode(image_file.read())

    if request is not None:
        return render_to_pdf_response(request, 'ncr/punchlistPDF_v2.html',
                                           {'pagesize': 'LETTER',
                                            'title': 'Reporte Punchlist',
                                            'resultados': resultados,
                                            'main_fotos': imagenes,
                                            'parque': parque,
                                            'titulo': titulo,
                                            'show_fotos': show_fotos,
                                            'img_solucionado': img_solucionado,
                                            'img_parcialsolucionado': img_parcialsolucionado,
                                            'img_nosolucionado': img_nosolucionado,
                                            'logo_saroen': logo_saroen,
                                            'colores' : colores,
                                            'estados': estados,
                                            'fecha': fecha,
                                            'nombre': nombre,
                                            }, content_type='application/pdf',
                                           response_class=HttpResponse)
    else:
        pdf = render_to_pdf('ncr/punchlistPDF_v2.html',
                            {'pagesize': 'LETTER',
                             'title': 'Reporte Punchlist',
                             'resultados': resultados,
                             'main_fotos': imagenes,
                             'parque': parque,
                             'titulo': titulo,
                             'show_fotos': show_fotos,
                             'img_solucionado': img_solucionado,
                             'img_parcialsolucionado': img_parcialsolucionado,
                             'img_nosolucionado': img_nosolucionado,
                             'logo_saroen': logo_saroen,
                             'colores': colores,
                             'estados': estados,
                             'fecha': fecha,
                             'nombre': nombre,
                             })
        return StringIO.StringIO(pdf)

def listFotos(resultados):
    main_fotos = {}
    for res in resultados:
        r=res.revision_set.all().order_by('-id')[0]
        results = Fotos.objects.filter(revision=r, principal=True).order_by('orden')
        if results.count() > 0:
            main_fotos[res.id] = results[0].reporte_img.url
    return main_fotos

def listFotos_v2(resultados):
    main_fotos = []
    resultados2 = resultados.order_by('componente__orden_punchlist','id')
    numero = 1
    for observacion in resultados2:
        r=observacion.revision_set.all().order_by('-id')[0]
        results = Fotos.objects.filter(revision=r, principal=True).order_by('orden')
        count = 1
        for f in results:
            cuadro_foto = {}
            cuadro_foto['numero'] = numero
            cuadro_foto['url'] = f.reporte_img.url
            cuadro_foto['texto'] = 'OBS_' + observacion.parque.codigo + '-' + observacion.aerogenerador.nombre + \
                                    '-' + str(observacion.observacion_id)
            if results.count() > 1:
                cuadro_foto['texto'] += ' (' + str(count) + ')'
                count += 1
            cuadro_foto['status'] = observacion.estado.nombre
            main_fotos.append(cuadro_foto)
        numero += 1
    return main_fotos

def punchlistResults(parque, aerogenerador, form):
    reparadas = form.cleaned_data['reparadas']
    resultados = Observacion.objects.filter(parque=parque, aerogenerador=aerogenerador, punchlist=True)
    resultados = resultados.exclude(cerrado=True)
    resultados = resultados.filter(fecha_observacion__gte=form.cleaned_data['fecha_from'],
                                   fecha_observacion__lte=form.cleaned_data['fecha_to'])
    if not reparadas:
        resultados = resultados.exclude(estado__nombre__exact='Solucionado').exclude(cerrado=True)
    main_fotos = listFotos_v2(resultados)
    if aerogenerador.nombre == u'General':
        titulo = 'LISTADO DE OBSERVACIONES GENERALES'
    elif aerogenerador.nombre == u'Puerto':
        titulo = 'LISTADO DE OBSERVACIONES EN PUERTO'
    else:
        titulo = 'LISTADO DE PENDIENTES AEROGENERADOR ' + aerogenerador.nombre
    resultados = resultados.order_by('componente__orden_punchlist','id')
    return [resultados, main_fotos, titulo]

def generateWord(parque, aerogenerador,form, incluir_fotos = True):
    [resultados, main_fotos, titulo] = punchlistResults(parque, aerogenerador, form)
    if parque.word:
        nombre_archivo = parque.word.file.name
    else:
        nombre_archivo = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) + '/ncr/punchlist.docx'
    f = open(nombre_archivo, 'rb')
    document = Document(f)
    document._body.clear_content()

    section = document.sections[0]
    section.page_height = Mm(279.4)
    section.page_width = Mm(215.9)
    h = document.add_heading(titulo, 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    for s in document.styles.latent_styles.element:
        logger.debug(s.name)
    table = document.add_table(rows=1, cols=4, style="Punchlist")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[0].width = Cm(0.8)
    hdr_cells[1].text = 'Item'
    hdr_cells[1].width = Cm(3.8)
    hdr_cells[2].text = 'Componente'
    hdr_cells[2].width = Cm(2.5)
    hdr_cells[3].text = u'Descripción'
    hdr_cells[3].width = Cm(8.75)
    # styles = document.styles
    # table.rows[0].style = "borderColor:red;background-color:blue"
    count = 1
    for r in resultados:
        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        row_cells[0].width = Cm(0.8)
        row_cells[1].text = 'OBS_' + parque.codigo + '-' + r.aerogenerador.nombre + '-' + str(r.observacion_id)
        row_cells[1].width = Cm(3.8)
        row_cells[2].text = r.componente.nombre
        row_cells[2].width = Cm(2.5)
        row_cells[3].text = r.nombre
        row_cells[3].width = Cm(8.75)
        count += 1
    if incluir_fotos:
        document.add_page_break()
        h = document.add_heading(u'Fotografías', 2)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=1, cols=2, style="Fotos")
        first = True
        idx = 0
        for r in main_fotos:
            if not first:
                if idx % 2 == 0:
                    row_cells = table.add_row().cells
                    celda = 0
                else:
                    celda = 1
            else:
                row_cells = table.rows[0].cells
                first = False
                celda = 0
            idx = idx + 1

            archivo = settings.BASE_DIR + r['url']
            c = row_cells[celda].paragraphs[0]
            aux = c.add_run()
            try:
                aux.add_picture(archivo, width=Cm(7.5))
                p = row_cells[celda].add_paragraph()
                p.text = r['texto']
                aux2 = p.add_run()
                if r['status'] == 'Solucionado':
                    status_img = os.path.join(settings.BASE_DIR, 'static/common/images/check-mark-3-64.gif')
                elif r['status'] == 'Parcialmente Solucionado':
                    status_img = os.path.join(settings.BASE_DIR, 'static/common/images/x-mark-64-amarillo.gif')
                elif r['status'] == 'No Solucionado':
                    status_img = os.path.join(settings.BASE_DIR, 'static/common/images/x-mark-64.gif')
                aux2.add_picture(status_img, width=Cm(0.4))
            except IOError:
                logger.debug('No se encontró imagen')
    target_stream = StringIO.StringIO()
    document.save(target_stream)
    return target_stream

@login_required(login_url='ingresar')
def punchlist(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Punchlist'
    contenido.subtitulo='Parque '+ parque.nombre
    contenido.menu = ['menu-ncr', 'menu2-punchlist']

    form = Punchlist(parque=parque)
    resultados = None

    show_fotos = False
    if request.method == 'POST':
        logger.debug('informePunchlist Post')
        form = Punchlist(request.POST, parque=parque)

        if form.is_valid():
            logger.debug("Form Valid")
            show_fotos = form.cleaned_data['fotos']
            if 'word' in request.POST:
                if not request.user.has_perm('usuarios.create_editables') :
                    raise PermissionDenied
                logger.debug('WORD')
                if len(form.cleaned_data['aerogenerador']) == 1:
                    ag = form.cleaned_data['aerogenerador'][0]
                    target_stream = generateWord(parque,ag,form,incluir_fotos=show_fotos)
                    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    fecha_str = form.cleaned_data['fecha'].strftime("%y%m%d")
                    nombre = 'PL_' + parque.codigo + '-' + ag.nombre + '_' + fecha_str + '.docx'
                    response['Content-Disposition'] = 'attachment; filename='+nombre
                    target_stream.flush()
                    ret_word = target_stream.getvalue()
                    target_stream.close()
                    response.write(ret_word)
                    return response
                elif len(form.cleaned_data['aerogenerador']) > 1:
                    response = HttpResponse(content_type='application/zip')
                    response['Content-Disposition'] = 'filename=punchlistWord.zip'
                    buff = StringIO.StringIO()
                    archive = zipfile.ZipFile(buff, 'w', zipfile.ZIP_DEFLATED)

                    for ag in form.cleaned_data['aerogenerador']:
                        target_stream = generateWord(parque, ag, form, incluir_fotos=show_fotos)
                        fecha_str = form.cleaned_data['fecha'].strftime("%y%m%d")
                        archivo_name = 'PL_' + parque.codigo + '-' + ag.nombre + '_' + fecha_str + '.docx'
                        logger.debug(archivo_name)
                        archive.writestr(archivo_name, target_stream.getvalue())
                    archive.close()
                    buff.flush()
                    ret_zip = buff.getvalue()
                    buff.close()
                    response.write(ret_zip)
                    return response
            else:
                if len(form.cleaned_data['aerogenerador']) == 1:
                    ag = form.cleaned_data['aerogenerador'][0]
                    [resultados, main_fotos, titulo] = punchlistResults(parque,ag,form)
                    logger.debug(resultados)
                    colores = form.cleaned_data['colores']
                    estados = form.cleaned_data['estados']
                    fecha_str = form.cleaned_data['fecha'].strftime("%y%m%d")
                    nombre = 'PL_' + parque.codigo + '-' + ag.nombre + '_' + fecha_str + '.pdf'
                    respuesta = generatePdf(parque,resultados,main_fotos,titulo,request,
                                            show_fotos=show_fotos,
                                            colores=colores,
                                            estados=estados,
                                            fecha=form.cleaned_data['fecha'],
                                            nombre=nombre)
                    respuesta['Content-Disposition'] = 'attachment; filename='+nombre
                    return respuesta
                elif len(form.cleaned_data['aerogenerador']) > 1:
                    response = HttpResponse(content_type='application/zip')
                    response['Content-Disposition'] = 'filename=punchlistPDF.zip'
                    buff = StringIO.StringIO()
                    archive = zipfile.ZipFile(buff, 'w', zipfile.ZIP_DEFLATED)

                    for ag in form.cleaned_data['aerogenerador']:
                        [resultados, main_fotos, titulo] = punchlistResults(parque, ag, form)
                        logger.debug(resultados)
                        colores = form.cleaned_data['colores']
                        estados = form.cleaned_data['estados']
                        fecha_str = form.cleaned_data['fecha'].strftime("%y%m%d")
                        archivo_name = 'PL_' + parque.codigo + '-' + ag.nombre + '_' + fecha_str + '.pdf'
                        archivo = generatePdf(parque, resultados, main_fotos, titulo,
                                              show_fotos=show_fotos,
                                              colores = colores,
                                              estados=estados,
                                              fecha=form.cleaned_data['fecha'],
                                              nombre=archivo_name)

                        logger.debug(archivo_name)
                        archive.writestr(archivo_name, archivo.getvalue())
                    archive.close()
                    buff.flush()
                    ret_zip = buff.getvalue()
                    buff.close()
                    response.write(ret_zip)
                    return response


    return TemplateResponse(request, 'ncr/punchlist.html',
        {'cont': contenido,
         'parque': parque,
         'aerogeneradores':aerogeneradores,
         'form': form,
         'resultados': resultados,
        })

@login_required(login_url='ingresar')
def observadores(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Listado de inspectores'
    contenido.subtitulo='Parque '+ parque.nombre
    contenido.menu = ['menu-principal', 'menu2-observadores']
    observadores = Observador.objects.all().order_by('id')
    response_data = {}

    # Me decido por interfaz RESTful. POST: Crear Nuevo, DELETE: Eliminar, PUT: Editar
    if request.method == 'POST':
        logger.debug('POST')
        editar = Observador.objects.get(id=int(request.POST['id']))
        editar.nombre = request.POST['nombre']
        editar.save()
        response_data['id'] = str(editar.id)
        response = json.dumps(response_data)
        return HttpResponse(
            response,
            content_type="application/json"
        )
    elif request.method == 'DELETE':
        #logger.debug('DELETE')
        data = dict(urlparse.parse_qsl(request.body))
        borrar = Observador.objects.get(id=data['id'])
        logger.debug("Borrando nombre=" +borrar.nombre)
        borrar.delete()
        return HttpResponse(
            '',
            content_type="application/json"
        )
    elif request.method == 'PUT':
        data=dict(urlparse.parse_qsl(request.body))
        nuevo = Observador(nombre = data['nombre'])
        nuevo.save()
        response_data['id'] = str(nuevo.id)
        response = json.dumps(response_data)
        return HttpResponse(
            response,
            content_type="application/json"
        )
    return TemplateResponse(request, 'ncr/agregarObservador.html',
        {'cont': contenido,
         'parque': parque,
         'observadores': observadores,
         'aerogeneradores':aerogeneradores,
        })

@login_required(login_url='ingresar')
def del_revision(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Listado de observadores'
    contenido.subtitulo='Parque '+ parque.nombre
    contenido.menu = ['menu-principal', 'menu2-observadores']
    observadores = Observador.objects.all().order_by('id')
    response_data = {}

    if request.method == 'POST':
        logger.debug('POST del_observacion')
        revision=Revision.objects.get(id=int(request.POST['del_id']))
        if not (request.user.has_perm('ncr.delete_revision') or request.user == revision.created_by) :
            raise PermissionDenied
        observacion = revision.observacion
        log_msg = "Se elimina revisión para parque " + parque.nombre + \
                  " - Aerogenerador - " + revision.observacion.aerogenerador.nombre + \
                  " - Observacion " + revision.observacion.nombre
        log = Log(texto=log_msg, tipo=3, user=request.user)
        log.save()
        revision.delete()
        observacion.save()  # Necesario para actualizar campos
        return HttpResponseRedirect(request.POST['back_url'])

@login_required(login_url='ingresar')
def close_observacion(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido=ContenidoContainer()
    contenido.user=request.user
    contenido.titulo=u'Listado de observadores'
    contenido.subtitulo='Parque '+ parque.nombre
    contenido.menu = ['menu-principal', 'menu2-observadores']
    observadores = Observador.objects.all().order_by('id')
    response_data = {}

    if request.method == 'POST':
        logger.debug('POST del_observacion')
        observacion=Observacion.objects.get(id=int(request.POST['del_id']))
        if not (request.user.has_perm('ncr.change_observacion') or request.user == observacion.created_by) :
            raise PermissionDenied
        if request.POST['cerrar'] == '1':
            observacion.msg_cerrado = request.POST['cierre_msg']
            observacion.cerrado = True
            log_msg = "Se cierra observacion para parque " + parque.nombre + \
                      " - Aerogenerador - " + observacion.aerogenerador.nombre + \
                      " - Observacion " + observacion.nombre
            log = Log(texto=log_msg, tipo=2, user=request.user)
            log.save()
        else:
            observacion.msg_cerrado = ''
            observacion.cerrado = False
            log_msg = "Se abre observacion para parque " + parque.nombre + \
                      " - Aerogenerador - " + observacion.aerogenerador.nombre + \
                      " - Observacion " + observacion.nombre
            log = Log(texto=log_msg, tipo=2, user=request.user)
            log.save()
        observacion.save()
        return HttpResponseRedirect(request.POST['back_url'])

@login_required(login_url='ingresar')
def imagenes_aerogenerador(request,slug,slug_ag):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    if slug_ag != 'resumen':
        aerogenerador = get_object_or_404(Aerogenerador, parque=parque, slug=slug_ag)

    if slug_ag == 'resumen':
        observaciones = Observacion.objects.filter(parque=parque)
    else:
        observaciones = Observacion.objects.filter(aerogenerador__idx__exact=aerogenerador.idx, parque=parque)

    grafico_estado = graficoBarrasSimple(observaciones, 'estado', EstadoRevision.objects.all().order_by('-id'),
                                         showall=True)

    observaciones = observaciones.exclude(estado__nombre__exact='Solucionado')

    grafico_componente = graficoBarrasSimple(observaciones, 'componente', Componente.objects.all())
    grafico_severidad = graficoBarrasSimple(observaciones, 'severidad', Severidad.objects.all(), showall=True)
    grafico_prioridad = graficoBarrasSimple(observaciones, 'prioridad', Prioridad.objects.all(), showall=True)
    grafico_subcomponente = graficoBarrasSimple(observaciones, 'sub_componente', Subcomponente.objects.all())
    grafico_tipo = graficoBarrasSimple(observaciones, 'tipo', Tipo.objects.all())
    if slug_ag == 'resumen':
        grafico_aerogenerador = graficoBarrasCompleto(observaciones, 'aerogenerador', aerogeneradores, showall=True)

    with open('static/common/js/ncr-charts.js', 'r') as myfile:
        data_js_file = myfile.read().replace('\n', '')

    parser = Parser()
    tree = parser.parse(data_js_file)
    node_estado = None
    node_columna = None
    node_graficofull2 = None
    data_found = False
    for node in nodevisitor.visit(tree):
        if isinstance(node,ast.Object) and node_estado is None:
            if hasattr(node,'properties'):
                node_estado = node
                logger.debug(node.to_ecma())
        if node_estado is not None and not data_found:
            if isinstance(node,ast.Assign):
                if node.left.value == 'data':
                    node.right.value = grafico_estado
                    data_found = True
        if data_found and node_columna is None:
            if isinstance(node, ast.Object):
                if hasattr(node, 'properties'):
                    node_columna = node
        if data_found and node_columna is not None:
            if isinstance(node,ast.Identifier):
                if node.value == 'GraficoFull2':
                    node_graficofull2 = node
        if data_found and node_graficofull2 is not None:
            if isinstance(node, ast.Object):
                if hasattr(node, 'properties'):
                    node_graficofull2 = node
                    break

    nodo_titulo = None
    nodo_data = None
    for node in nodevisitor.visit(node_columna):
        if isinstance(node,ast.Assign) and nodo_titulo is None:
            if node.left.value == 'text':
                node.right.value = "'Número de observaciones por componente'"
                nodo_titulo = node
        if isinstance(node, ast.Assign) and nodo_titulo is not None:
            if node.left.value == 'data':
                node.right.value = grafico_componente
                nodo_data = node
                break

    grafico_estado_b64 = ''
    if node_estado is not None:
        r = requests.post('http://67.205.142.111:8001', data={'infile': node_estado.to_ecma(), 'b64': 'true', 'scale':2})
        if r.status_code == 200:
            grafico_estado_b64 = r.text
    grafico_componente_b64 = ''
    if node_columna is not None:
        r = requests.post('http://67.205.142.111:8001', data={'infile': node_columna.to_ecma(), 'b64': 'true', 'scale':2})
        if r.status_code == 200:
            grafico_componente_b64 = r.text

    grafico_subcomponente_b64 = ''
    nodo_titulo.right.value = "'Número de observaciones por sub-componente'"
    nodo_data.right.value = grafico_subcomponente
    r = requests.post('http://67.205.142.111:8001', data={'infile': node_columna.to_ecma(), 'b64': 'true', 'scale': 2})
    if r.status_code == 200:
        grafico_subcomponente_b64 = r.text

    grafico_tipo_b64 = ''
    nodo_titulo.right.value = "'Número de observaciones por tipo'"
    nodo_data.right.value = grafico_tipo
    r = requests.post('http://67.205.142.111:8001', data={'infile': node_columna.to_ecma(), 'b64': 'true', 'scale': 2})
    if r.status_code == 200:
        grafico_tipo_b64 = r.text

    grafico_severidad_b64 = ''
    nodo_titulo.right.value = "'Número de observaciones por severidad'"
    nodo_data.right.value = grafico_severidad
    r = requests.post('http://67.205.142.111:8001', data={'infile': node_columna.to_ecma(), 'b64': 'true', 'scale': 2})
    if r.status_code == 200:
        grafico_severidad_b64 = r.text

    grafico_prioridad_b64 = ''
    nodo_titulo.right.value = "'Número de observaciones por prioridad'"
    nodo_data.right.value = grafico_prioridad
    r = requests.post('http://67.205.142.111:8001', data={'infile': node_columna.to_ecma(), 'b64': 'true', 'scale': 2})
    if r.status_code == 200:
        grafico_prioridad_b64 = r.text

    grafico_resumen_b64 = None
    nodo_titulo = None
    if slug_ag == 'resumen':
        for node in nodevisitor.visit(node_graficofull2):
            if isinstance(node, ast.Assign) and nodo_titulo is None:
                if node.left.value == 'text':
                    node.right.value = "'Severidad por aerogenerador'"
                    nodo_titulo = node
            if isinstance(node, ast.Assign) and nodo_titulo is not None:
                if node.left.value == 'series':
                    node.right.value = grafico_aerogenerador
                if node.left.value == 'fontSize':
                    node.right.value = "'4px'"
                    break
        logger.debug(node_graficofull2.to_ecma())
        r = requests.post('http://67.205.142.111:8001', data={'infile': node_graficofull2.to_ecma(), 'b64': 'true', 'scale': 4})
        if r.status_code == 200:
            grafico_resumen_b64 = r.text

    if slug_ag == 'resumen':
        titulo = 'Gráficas Resumen'
        nombre_archivo = 'graficas_resumen.pdf'
    else:
        titulo = 'Gráficas ' + aerogenerador.nombre
        nombre_archivo = 'graficas_' + aerogenerador.slug + '.pdf'

    with open(os.path.join(settings.BASE_DIR, 'static/common/images/saroenlogo.png'), "rb") as image_file:
        logo_saroen = base64.b64encode(image_file.read())

    fecha = datetime.now()
    return render_to_pdf_response(request, 'ncr/graficos_pdf.html',
                                  {'pagesize': 'LETTER',
                                   'title': 'Graficas Aerogenerador',
                                   'parque': parque,
                                   'titulo': titulo,
                                   'logo_saroen': logo_saroen,
                                   'grafico_estado_b64': grafico_estado_b64,
                                   'grafico_componente_b64': grafico_componente_b64,
                                   'grafico_subcomponente_b64': grafico_subcomponente_b64,
                                   'grafico_tipo_b64': grafico_tipo_b64,
                                   'grafico_severidad_b64': grafico_severidad_b64,
                                   'grafico_prioridad_b64': grafico_prioridad_b64,
                                   'grafico_resumen_b64': grafico_resumen_b64,
                                   'fecha': fecha,
                                   #'nombre': nombre,
                                   }, content_type='application/pdf',
                                  filename= nombre_archivo,
                                  response_class=HttpResponse)

@login_required(login_url='ingresar')
def configuracion(request,slug):
    parque = get_object_or_404(ParqueSolar, slug=slug)
    aerogeneradores = Aerogenerador.objects.filter(parque=parque).order_by('idx')
    contenido = ContenidoContainer()
    contenido.user = request.user
    contenido.titulo = u'Edición'
    contenido.subtitulo = 'Componentes, Sub-componentes y Tipo'
    contenido.menu = ['menu-ncr', 'menu2-configuracion-ncr']

    if request.method == 'POST':
        logger.debug('ncr-configuracion')
        if 'orden' in request.POST:
            target = request.POST['target']
            orden = json.loads(request.POST.get('orden'))
            orden_actividad = 1
            for obj in orden:
                id = obj['id']
                c = None
                if target == 'componente':
                    c = get_object_or_404(Componente, id=id)
                elif target == 'subcomponente':
                    c = get_object_or_404(Subcomponente,id=id)
                elif target == 'tipo':
                    c = get_object_or_404(Tipo, id=id)

                if c is not None:
                    c.orden_punchlist = orden_actividad
                    c.save()
                orden_actividad += 1

        elif 'agregar' in request.POST:
            nombre = request.POST['nombre']
            target = request.POST['target']
            c = None
            if target == 'componente':
                c = Componente(nombre=nombre)
            elif target == 'subcomponente':
                c = Subcomponente(nombre=nombre)
            elif target == 'tipo':
                c = Tipo(nombre=nombre)

            if c is not None:
                try:
                    c.save()
                    messages.add_message(request, messages.SUCCESS, target.title() + ' agregado con éxito')
                except IntegrityError:
                    messages.add_message(request, messages.ERROR, target.title() + ' ya existe')

        elif 'editar' in request.POST:
            # Falta un try para nombres duplicados (mostrar error)
            id = int(request.POST['id'])
            nombre = request.POST['nombre']
            target = request.POST['target']
            c = None
            if target == 'componente':
                c = get_object_or_404(Componente, id=id)
            elif target == 'subcomponente':
                c = get_object_or_404(Subcomponente, id=id)
            elif target == 'tipo':
                c = get_object_or_404(Tipo, id=id)

            if c is not None:
                c.nombre = nombre
                try:
                    c.save()
                    messages.add_message(request, messages.SUCCESS, target.title() + ' editado con éxito')
                except IntegrityError:
                    messages.add_message(request, messages.ERROR, target.title() + ' ya existe')

        elif 'eliminar' in request.POST:
            id = int(request.POST['id'])
            target = request.POST['target']
            c = None
            if target == 'componente':
                c = get_object_or_404(Componente, id=id)
            elif target == 'subcomponente':
                c = get_object_or_404(Subcomponente, id=id)
            elif target == 'tipo':
                c = get_object_or_404(Tipo, id=id)

            if c is not None:
                c.delete()
                messages.add_message(request, messages.SUCCESS, target.title() + ' eliminado con éxito')
    componentes = Componente.objects.all().order_by('orden_punchlist')
    subcomponentes = Subcomponente.objects.all().order_by('orden_punchlist')
    tipos = Tipo.objects.all().order_by('orden_punchlist')

    return TemplateResponse(request, 'ncr/configuracion.html',
                            {'cont': contenido,
                             'parque': parque,
                             'aerogeneradores': aerogeneradores,
                             'componentes': componentes,
                             'subcomponentes': subcomponentes,
                             'tipos': tipos})